#!/usr/bin/env python
"""DOCX to Markdown converter module.

This module provides functions to convert Microsoft Word DOCX files to Markdown format.
"""

from pathlib import Path
from typing import TYPE_CHECKING, Any

from aichemist_transmutation_codex.config import ConfigManager, LogManager

# Setup logger
log_manager = LogManager()
logger = log_manager.get_converter_logger("docx2md")

# Try to import necessary libraries
# python-docx for DOCX parsing, mammoth as an alternative
try:
    import docx
    from docx import oxml
    from docx.oxml.ns import nspfxmap, qn  # Import nspfxmap and qn directly

    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not found. Install with 'pip install python-docx'")
    DOCX_AVAILABLE = False

try:
    import mammoth

    MAMMOTH_AVAILABLE = True
except ImportError:
    logger.warning("mammoth not found. Install with 'pip install mammoth'")
    MAMMOTH_AVAILABLE = False

if TYPE_CHECKING:
    from docx.table import Table as DocxTable  # type: ignore
    from docx.text.paragraph import Paragraph as DocxParagraph  # type: ignore


def convert_docx_to_md(
    input_path: str | Path,
    output_path: str | Path | None = None,
    **kwargs: Any,
) -> Path:
    """Convert a DOCX file to Markdown format.

    This function converts a Microsoft Word document (DOCX) to Markdown format,
    preserving formatting, headers, lists, tables, and images where possible.

    Args:
        input_path: Path to the input DOCX file
        output_path: Path for the output Markdown file (optional)
        **kwargs: Additional arguments:
            - image_dir: Directory to save extracted images (default: same as output)
            - extract_images: Whether to extract and include images (default: True)
            - extract_tables: Whether to extract and format tables (default: True)
            - use_mammoth: Force using mammoth library instead of python-docx (default: False)

    Returns:
        Path to the generated Markdown file

    Raises:
        FileNotFoundError: If input file doesn't exist
        ValueError: If input file is not a DOCX file
        ImportError: If required dependencies are not installed
    """  # noqa: E501
    # Get config
    config = ConfigManager()
    settings = config.get_converter_config("docx2md")

    # Check dependencies
    if not DOCX_AVAILABLE and not MAMMOTH_AVAILABLE:
        logger.error(
            "Neither python-docx nor mammoth is available for DOCX conversion."
        )
        raise ImportError("python-docx or mammoth is required.")

    input_path = Path(input_path).resolve()
    if not input_path.exists():
        logger.error(f"Input DOCX file not found: {input_path}")
        raise FileNotFoundError(f"Input file not found: {input_path}")
    if input_path.suffix.lower() not in (".docx", ".doc"):  # Allow .doc as well
        logger.error(f"Invalid input file type: {input_path.suffix}")
        raise ValueError(f"Input file must be a DOCX file, got: {input_path}")

    output_path = (
        Path(output_path).resolve() if output_path else input_path.with_suffix(".md")
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Extract options from kwargs or config
    extract_images = kwargs.get("extract_images", settings.get("extract_images", True))
    extract_tables = kwargs.get("extract_tables", settings.get("extract_tables", True))
    use_mammoth = (
        kwargs.get("use_mammoth", settings.get("use_mammoth", False))
        or not DOCX_AVAILABLE
    )
    # Handle image directory relative to output path
    img_dir_setting = kwargs.get("image_dir", settings.get("image_dir", "images"))
    image_dir = output_path.parent / img_dir_setting

    if extract_images:
        image_dir.mkdir(parents=True, exist_ok=True)
        logger.debug(f"Image directory set to: {image_dir}")

    logger.info(f"Converting {input_path} to Markdown...")

    try:
        if use_mammoth and MAMMOTH_AVAILABLE:
            logger.info("Using mammoth engine for conversion.")
            return _convert_with_mammoth(
                input_path, output_path, image_dir, extract_images
            )
        elif DOCX_AVAILABLE:
            logger.info("Using python-docx engine for conversion.")
            return _convert_with_python_docx(
                input_path, output_path, image_dir, extract_images, extract_tables
            )
        else:
            # Fallback if preferred engine fails or isn't available
            logger.error("No suitable engine found for DOCX conversion.")
            raise RuntimeError("Could not find a suitable DOCX conversion engine.")
    except Exception as e:
        logger.exception(f"Error during DOCX to Markdown conversion: {e}")
        raise RuntimeError(f"Error converting DOCX to Markdown: {e}") from e


def _convert_with_mammoth(
    input_path: Path,
    output_path: Path,
    image_dir: Path,
    extract_images: bool,
) -> Path:
    """Convert DOCX to Markdown using the mammoth library."""
    if not MAMMOTH_AVAILABLE:
        logger.error("Mammoth library called but not available.")
        raise ImportError("mammoth library is not available.")

    options = {}
    if extract_images:
        image_dir.mkdir(parents=True, exist_ok=True)

        def image_handler(image):
            # Use content type to determine extension
            ext = image.content_type.split("/")[1] if image.content_type else "png"
            # Create a more robust unique filename
            img_hash = hash(image.byte_array)
            image_filename = f"image_{input_path.stem}_{img_hash}.{ext}"
            image_path = image_dir / image_filename

            try:
                with open(image_path, "wb") as f:
                    f.write(image.byte_array)

                # Relative path from output MD file to image
                relative_path = image_path.relative_to(output_path.parent).as_posix()
                logger.debug(
                    f"Extracted image to {image_path} (relative: {relative_path})"
                )
                return {"src": relative_path}
            except Exception as img_err:
                logger.error(f"Failed to save image {image_filename}: {img_err}")
                return {"src": ""}  # Return empty src on failure

        options["convert_image"] = mammoth.images.img_element(image_handler)

    try:
        with open(input_path, "rb") as docx_file:
            result = mammoth.convert_to_markdown(docx_file, **options)

        with open(output_path, "w", encoding="utf-8") as md_file:
            md_file.write(result.value)

        if result.messages:
            for message in result.messages:
                logger.warning(f"Mammoth conversion message: {message}")

        logger.info(f"Successfully converted to {output_path} using mammoth")
        return output_path
    except Exception as e:
        logger.exception(f"Mammoth conversion failed: {e}")
        raise RuntimeError(f"Mammoth conversion failed: {e}") from e


def _convert_with_python_docx(
    input_path: Path,
    output_path: Path,
    image_dir: Path,
    extract_images: bool,
    extract_tables: bool,
) -> Path:
    """Convert DOCX to Markdown using the python-docx library."""
    if not DOCX_AVAILABLE:
        logger.error("python-docx library called but not available.")
        raise ImportError("python-docx library is not available.")

    try:
        # Pass the path string to Document
        doc = docx.Document(str(input_path))
        md_content = []
        image_counter = 0  # For unique image naming

        for element in doc.element.body:
            if isinstance(element, docx.oxml.text.paragraph.CT_P):  # type: ignore
                para = docx.text.paragraph.Paragraph(element, doc)  # type: ignore
                if not para.text.strip():
                    md_content.append("")
                    continue

                # Handle headings (check style safely)
                style_name = getattr(para.style, "name", "")
                if style_name.startswith("Heading"):
                    level = int(style_name[-1]) if style_name[-1].isdigit() else 1
                    md_content.append(f"\n{'#' * level} {para.text.strip()}\n")
                # Handle list items
                elif (
                    style_name.startswith("List Paragraph")
                    or para.style.name == "List Bullet"
                ):
                    # Basic list handling (needs improvement for levels)
                    md_content.append(f"* {para.text.strip()}")
                else:
                    formatted_text = _format_paragraph(para)
                    md_content.append(formatted_text)

                # Extract images within paragraphs
                if extract_images:
                    for run in para.runs:
                        for inline_shape in run.element.findall(
                            ".//wp:inline",
                            namespaces=nspfxmap,  # Use nspfxmap directly
                        ):  # type: ignore
                            blip = inline_shape.find(
                                ".//a:blip",
                                namespaces=nspfxmap,  # Use nspfxmap directly
                            )  # type: ignore
                            if blip is not None:
                                embed_id = blip.get(qn("r:embed"))  # Use qn directly
                                if embed_id:
                                    image_part = doc.part.related_parts[embed_id]
                                    image_bytes = image_part.blob
                                    ext = image_part.content_type.split("/")[1]
                                    image_counter += 1
                                    image_filename = (
                                        f"image_{input_path.stem}_{image_counter}.{ext}"
                                    )
                                    image_path = image_dir / image_filename
                                    try:
                                        with open(image_path, "wb") as img_file:
                                            img_file.write(image_bytes)
                                        relative_path = image_path.relative_to(
                                            output_path.parent
                                        ).as_posix()
                                        md_content.append(
                                            f"![{image_filename}]({relative_path})"
                                        )
                                        logger.debug(
                                            f"Extracted image {image_filename}"
                                        )
                                    except Exception as img_err:
                                        logger.error(
                                            f"Failed to save image {image_filename}: {img_err}"
                                        )

            elif extract_tables and isinstance(
                element, oxml.table.CT_Tbl
            ):  # Keep oxml for this
                table = docx.table.Table(element, doc)  # type: ignore
                md_table = _format_table(table)
                if md_table:
                    md_content.append("\n" + md_table + "\n")

        with open(output_path, "w", encoding="utf-8") as md_file:
            # Join paragraphs, handling potential multiple newlines correctly
            output_text = "\n".join(md_content).replace("\n\n\n", "\n\n")
            md_file.write(output_text)

        logger.info(f"Successfully converted to {output_path} using python-docx")
        return output_path
    except Exception as e:
        logger.exception(f"python-docx conversion failed: {e}")
        raise RuntimeError(f"python-docx conversion failed: {e}") from e


def _format_paragraph(paragraph: "DocxParagraph") -> str:
    """Formats a python-docx Paragraph object into a Markdown string.

    Iterates through the runs in a paragraph, applying bold and italic
    Markdown formatting based on the run's properties. Underline is converted
    to italic as Markdown lacks native underline support.

    Args:
        paragraph (DocxParagraph): The `docx.text.paragraph.Paragraph` object to format.

    Returns:
        str: A string containing the Markdown representation of the paragraph.
    """
    formatted_text = ""
    for run in paragraph.runs:
        text = run.text
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        if (
            run.underline
        ):  # Convert underline to italic as Markdown doesn't have native underline
            text = f"*{text}*"
        # Add more formatting like strikethrough, highlight if needed
        formatted_text += text
    return formatted_text


def _format_table(table: "DocxTable") -> str | None:
    """Converts a python-docx Table object into a Markdown table string.

    Constructs a Markdown table with a header row and data rows. Cells containing
    newlines are converted to use `<br>` tags for line breaks within the table cell.
    Rows with mismatched column counts compared to the header are skipped with a warning.

    Args:
        table (DocxTable): The `docx.table.Table` object to format.

    Returns:
        str | None: A string containing the Markdown table, or None if the input
            table has no rows. Returns a placeholder string "*Error processing table*"
            if an exception occurs during formatting.
    """
    if not table.rows:
        return None

    md_table = []
    try:
        # Header row
        header_cells = table.rows[0].cells
        header_row = [
            " ".join(p.text for p in cell.paragraphs).strip() for cell in header_cells
        ]
        md_table.append("| " + " | ".join(header_row) + " |")
        md_table.append("| " + " | ".join(["---"] * len(header_row)) + " |")

        # Data rows
        for row in table.rows[1:]:
            row_cells = row.cells
            # Ensure row has same number of cells as header for valid Markdown
            if len(row_cells) == len(header_row):
                row_data = [
                    " ".join(p.text for p in cell.paragraphs)
                    .strip()
                    .replace("\n", "<br>")
                    for cell in row_cells
                ]
                md_table.append("| " + " | ".join(row_data) + " |")
            else:
                logger.warning(
                    f"Skipping table row with mismatched column count ({len(row_cells)} vs {len(header_row)})"
                )

        return "\n".join(md_table)
    except Exception as table_err:
        logger.error(f"Failed to format table: {table_err}")
        return "*Error processing table*"


# Define an alias for backward compatibility
docx_to_md = convert_docx_to_md
