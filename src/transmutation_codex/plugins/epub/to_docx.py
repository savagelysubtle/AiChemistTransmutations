"""EPUB to DOCX converter.

This module provides functionality to convert EPUB files to Microsoft Word DOCX format.
Extracts content from EPUB and generates a well-formatted DOCX document.
"""

from pathlib import Path
from typing import Any

try:
    import ebooklib
    from ebooklib import epub

    EBOOKLIB_AVAILABLE = True
except ImportError:
    EBOOKLIB_AVAILABLE = False

try:
    from bs4 import BeautifulSoup

    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from transmutation_codex.core import (
    check_feature_access,
    check_file_size_limit,
    complete_operation,
    get_log_manager,
    publish,
    record_conversion_attempt,
    start_operation,
    update_progress,
)
from transmutation_codex.core.decorators import converter
from transmutation_codex.core.events import ConversionEvent, EventTypes
from transmutation_codex.core.exceptions import raise_conversion_error

# Setup logger
logger = get_log_manager().get_converter_logger("epub2docx")


@converter(
    source_format="epub",
    target_format="docx",
    description="Convert EPUB to Microsoft Word DOCX format",
    required_dependencies=["ebooklib", "bs4", "docx"],
    priority=10,
    version="1.0.0",
)
def convert_epub_to_docx(
    input_path: str | Path,
    output_path: str | Path | None = None,
    **kwargs: Any,
) -> Path:
    """Convert EPUB file to DOCX format.

    This function converts EPUB files to Microsoft Word DOCX format while
    preserving text content and basic formatting.

    Args:
        input_path: Path to input EPUB file
        output_path: Path for output DOCX file (auto-generated if None)
        **kwargs: Additional options:
            - `include_toc` (bool): Whether to include table of contents.
                                   Defaults to True.
            - `chapter_breaks` (bool): Whether to add page breaks between chapters.
                                     Defaults to True.
            - `preserve_formatting` (bool): Whether to preserve basic formatting.
                                          Defaults to True.
            - `include_metadata` (bool): Whether to include EPUB metadata.
                                      Defaults to True.
            - `font_name` (str): Font name for the document.
                               Defaults to "Calibri".
            - `font_size` (int): Font size in points.
                               Defaults to 11.

    Returns:
        Path: The path to the generated DOCX file.

    Raises:
        ValidationError: If input or output paths are invalid, or dependencies are missing.
        ConversionError: If the conversion process fails.
    """
    logger.info(f"Attempting to convert EPUB to DOCX: {input_path}")

    # Validate dependencies
    if not EBOOKLIB_AVAILABLE:
        raise_conversion_error("ebooklib is required for EPUB conversion")
    if not BS4_AVAILABLE:
        raise_conversion_error("beautifulsoup4 is required for HTML parsing")
    if not DOCX_AVAILABLE:
        raise_conversion_error("python-docx is required for DOCX generation")

    # Start operation
    operation = start_operation(
        "conversion",
        total_steps=100,
        description=f"Converting EPUB to DOCX: {Path(input_path).name}",
    )

    try:
        # Check licensing and file size
        check_feature_access("epub2docx")
        check_file_size_limit(input_path, max_size_mb=100)
        record_conversion_attempt("epub2docx")

        # Convert paths
        input_path = Path(input_path)
        if output_path is None:
            output_path = input_path.with_suffix(".docx")
        else:
            output_path = Path(output_path)

        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        logger.info(f"Converting EPUB to DOCX: {input_path} -> {output_path}")

        # Parse options
        include_toc = kwargs.get("include_toc", True)
        chapter_breaks = kwargs.get("chapter_breaks", True)
        preserve_formatting = kwargs.get("preserve_formatting", True)
        include_metadata = kwargs.get("include_metadata", True)
        font_name = kwargs.get("font_name", "Calibri")
        font_size = kwargs.get("font_size", 11)

        update_progress(operation, 10, "Loading EPUB file...")

        # Load EPUB file
        try:
            book = epub.read_epub(str(input_path))
        except Exception as e:
            raise_conversion_error(f"Failed to load EPUB file: {e}")

        logger.info(f"EPUB loaded: {book.get_metadata('DC', 'title')}")

        update_progress(operation, 20, "Creating DOCX document...")

        # Create DOCX document
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)

        update_progress(operation, 30, "Processing EPUB metadata...")

        # Add metadata if requested
        if include_metadata:
            title = book.get_metadata("DC", "title")
            if title and title[0]:
                heading = doc.add_heading(title[0][0], level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            author = book.get_metadata("DC", "creator")
            if author and author[0]:
                author_para = doc.add_paragraph(f"by {author[0][0]}")
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                author_para.runs[0].font.size = Pt(font_size + 2)
                doc.add_paragraph()  # Spacer

            description = book.get_metadata("DC", "description")
            if description and description[0]:
                desc_para = doc.add_paragraph(f"{description[0][0]}")
                desc_para.runs[0].font.italic = True
                doc.add_paragraph()  # Spacer

        # Add table of contents if requested
        if include_toc:
            doc.add_heading("Table of Contents", level=1)

            # Get spine items
            spine_items = book.spine
            for i, (item_id, _) in enumerate(spine_items):
                item = book.get_item_by_id(item_id)
                if item and item.get_name():
                    toc_para = doc.add_paragraph(
                        f"{i + 1}. {item.get_name()}", style="List Number"
                    )

            doc.add_page_break()

        update_progress(operation, 40, "Converting chapters...")

        # Process chapters
        spine_items = book.spine
        total_items = len(spine_items)

        for i, (item_id, _) in enumerate(spine_items):
            logger.info(f"Processing chapter {i + 1}/{total_items}")
            update_progress(
                operation,
                40 + (i / total_items) * 50,
                f"Processing chapter {i + 1}",
            )

            item = book.get_item_by_id(item_id)
            if not item:
                continue

            # Add chapter title
            if item.get_name():
                doc.add_heading(item.get_name(), level=1)

            # Process content
            content = item.get_content()
            if content:
                # Parse HTML content
                soup = BeautifulSoup(content, "html.parser")

                # Convert to DOCX
                _html_to_docx(doc, soup, preserve_formatting, font_name, font_size)

            # Add page break between chapters
            if chapter_breaks and i < total_items - 1:
                doc.add_page_break()

        update_progress(operation, 90, "Saving DOCX file...")

        # Save DOCX file
        doc.save(str(output_path))

        # Publish success event
        publish(
            ConversionEvent(
                event_type=EventTypes.CONVERSION_COMPLETED,
                input_file=str(input_path),
                output_file=str(output_path),
                conversion_type="epub2docx",
            )
        )

        complete_operation(operation, {"output_path": str(output_path)})
        logger.info(f"EPUB to DOCX conversion completed: {output_path}")

        return output_path

    except Exception as e:
        logger.exception(f"EPUB to DOCX conversion failed: {e}")
        publish(
            ConversionEvent(
                event_type=EventTypes.CONVERSION_FAILED,
                input_file=str(input_path),
                conversion_type="epub2docx",
            )
        )
        raise_conversion_error(f"EPUB to DOCX conversion failed: {e}")


def _html_to_docx(
    doc: "Document",
    soup: "BeautifulSoup",
    preserve_formatting: bool,
    font_name: str,
    font_size: int,
) -> None:
    """Convert HTML content to DOCX paragraphs and add to document.

    Args:
        doc: The python-docx Document object
        soup: BeautifulSoup object containing HTML content
        preserve_formatting: Whether to preserve basic formatting
        font_name: Font name for the document
        font_size: Font size in points
    """
    for element in soup.find_all(
        [
            "p",
            "div",
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
            "ul",
            "ol",
            "li",
            "blockquote",
        ]
    ):
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            # Headings
            level = int(element.name[1])
            text = element.get_text().strip()
            if text:
                doc.add_heading(text, level=level)

        elif element.name == "p":
            # Paragraphs
            text = element.get_text().strip()
            if text:
                para = doc.add_paragraph()
                if preserve_formatting:
                    _add_formatted_text(para, element, font_name, font_size)
                else:
                    run = para.add_run(text)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

        elif element.name == "blockquote":
            # Blockquotes
            text = element.get_text().strip()
            if text:
                para = doc.add_paragraph(text, style="Intense Quote")

        elif element.name in ["ul", "ol"]:
            # Lists
            list_items = element.find_all("li", recursive=False)
            for li in list_items:
                text = li.get_text().strip()
                if text:
                    style = "List Bullet" if element.name == "ul" else "List Number"
                    para = doc.add_paragraph(text, style=style)

        elif element.name == "li":
            # List items (handled by parent ul/ol)
            continue

        elif element.name == "div":
            # Divs - just get text content
            text = element.get_text().strip()
            if text:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.name = font_name
                run.font.size = Pt(font_size)


def _add_formatted_text(
    para: "Paragraph", element: "BeautifulSoup", font_name: str, font_size: int
) -> None:
    """Add formatted text to a paragraph preserving basic HTML formatting.

    Args:
        para: The python-docx Paragraph object
        element: BeautifulSoup element containing text
        font_name: Font name for the document
        font_size: Font size in points
    """
    # Check for bold/italic formatting
    for child in element.descendants:
        if isinstance(child, str):
            text = child.strip()
            if text:
                run = para.add_run(text)
                run.font.name = font_name
                run.font.size = Pt(font_size)

                # Check if parent elements have formatting
                parent = child.parent
                while parent:
                    if parent.name in ["strong", "b"]:
                        run.font.bold = True
                    if parent.name in ["em", "i"]:
                        run.font.italic = True
                    if parent.name in ["u"]:
                        run.font.underline = True
                    if parent == element:
                        break
                    parent = parent.parent
