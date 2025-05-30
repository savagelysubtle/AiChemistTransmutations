import logging
import sys
from collections.abc import Callable
from pathlib import Path
from typing import Any

# Try to import the main conversion function from the docx2pdf library.
# docx2pdf is used to convert DOCX files to PDF format.
# It typically relies on Microsoft Word (on Windows) or LibreOffice (on Linux/macOS)
# being installed on the system.
try:
    from docx2pdf import convert as docx2pdf_convert

    # Check if the underlying OS-specific conversion tools are available
    # This isn't perfectly checkable by docx2pdf itself before a conversion attempt in all cases,
    # but we can make a note. The library handles errors gracefully if tools are missing.
    _docx2pdf_available = True
except ImportError as e:
    # If docx2pdf is not installed, set a flag and store the error.
    # The actual ImportError will be raised if a conversion is attempted.
    docx2pdf_convert = None  # type: ignore
    _docx2pdf_available = False
    _docx2pdf_import_error = e

# Configure a logger for this module.
# This allows recording of information, warnings, and errors.
logger = logging.getLogger(__name__)


def convert_docx_to_pdf(
    input_path: str | Path,
    output_path: str | Path | None = None,
    **kwargs: Any,
) -> Path:
    """Converts a DOCX file to a PDF document using the docx2pdf library.

    This function takes a DOCX (Microsoft Word) file and attempts to convert it
    to a PDF document. The success of this conversion often depends on having
    Microsoft Word (on Windows) or LibreOffice (on other platforms) installed
    and accessible by the `docx2pdf` library.

    Args:
        input_path (Union[str, Path]): The path to the input DOCX file.
        output_path (Union[str, Path, None], optional): The desired path for the
            output PDF file. If None, the output PDF will be saved in the same
            directory as the input DOCX file, with the same base name but a
            '.pdf' extension. Defaults to None.
        **kwargs (Any): Additional keyword arguments. Currently, 'progress_callback'
            is recognized:
            progress_callback (Callable[[int, str], None]): A function to call
                with progress updates (percentage, message). `docx2pdf` itself
                does not offer fine-grained progress, so this will be simulated.

    Returns:
        Path: The absolute path to the generated PDF file.

    Raises:
        ImportError: If the `docx2pdf` library is not installed.
        FileNotFoundError: If the input DOCX file (`input_path`) is not found.
        RuntimeError: If the DOCX to PDF conversion process fails for any reason
                      (e.g., Word/LibreOffice not found or inaccessible, issues
                       with the document itself).
        TypeError: If `input_path` or `output_path` (if provided) are not
                   valid path types (str or Path).
    """
    # Retrieve the progress_callback from kwargs, if it exists.
    progress_callback: Callable[[int, str], None] | None = kwargs.get(
        "progress_callback"
    )

    # Check if docx2pdf was imported successfully.
    if not _docx2pdf_available or docx2pdf_convert is None:
        logger.error(
            f"docx2pdf library is not installed or available. Original error: {_docx2pdf_import_error}"
        )
        raise ImportError(
            "The `docx2pdf` library is required for DOCX to PDF conversion but is not installed "
            f"or its dependencies are missing. Please install it (e.g., 'uv pip install docx2pdf'). "
            f"Original error: {_docx2pdf_import_error}"
        ) from _docx2pdf_import_error

    # Validate and resolve the input_path.
    if not isinstance(input_path, (str, Path)):
        err_msg = f"input_path must be a string or Path object. Got: {type(input_path)}"
        logger.error(err_msg)
        raise TypeError(err_msg)
    input_file_path = Path(input_path).resolve()  # Get absolute path

    if not input_file_path.exists() or not input_file_path.is_file():
        err_msg = f"Input DOCX file not found: {input_file_path}"
        logger.error(err_msg)
        raise FileNotFoundError(err_msg)

    if input_file_path.suffix.lower() not in [".docx"]:
        logger.warning(
            f"Input file '{input_file_path.name}' does not have a .docx extension. "
            "Conversion will be attempted, but it might fail if the file is not a valid DOCX."
        )

    # Determine and prepare the output_path.
    if output_path is None:
        # If no output_path is given, create one next to the input file
        # with the same name but a .pdf extension.
        output_file_path = input_file_path.with_suffix(".pdf")
    else:
        if not isinstance(output_path, (str, Path)):
            err_msg = f"output_path must be a string or Path object if provided. Got: {type(output_path)}"
            logger.error(err_msg)
            raise TypeError(err_msg)
        output_file_path = Path(output_path).resolve()

    # Ensure the directory for the output file exists.
    # The parent directory of output_file_path is created if it doesn't exist.
    output_file_path.parent.mkdir(parents=True, exist_ok=True)

    logger.info(
        f"Starting DOCX to PDF conversion: '{input_file_path}' -> '{output_file_path}'"
    )

    if progress_callback:
        # Simulate initial progress as docx2pdf is a single-step conversion.
        progress_callback(0, "Initializing conversion...")
        progress_callback(10, "Preparing to convert DOCX to PDF...")

    try:
        # Perform the conversion.
        # The `docx2pdf.convert()` function takes the input DOCX path
        # and optionally the output PDF path.
        # If output_file_path is not provided to docx2pdf.convert, it creates
        # the PDF in the same directory as the input.
        # We are providing it explicitly.
        if progress_callback:
            progress_callback(30, "Conversion in progress...")

        docx2pdf_convert(str(input_file_path), str(output_file_path))

        if progress_callback:
            # Simulate completion.
            progress_callback(100, "Conversion successful.")

        logger.info(
            f"Successfully converted '{input_file_path}' to '{output_file_path}'"
        )
        return output_file_path.resolve()  # Return the absolute path of the output.

    except Exception as e:
        # Catch a broad range of exceptions that docx2pdf might raise.
        # This can include errors if Word/LibreOffice is not found,
        # if there are permission issues, or if the document is corrupt.
        # The library itself might raise specific errors, but a general Exception
        # catch is safer here.
        error_message = (
            f"DOCX to PDF conversion failed for '{input_file_path}'. Error: {e}"
        )
        logger.error(error_message, exc_info=True)  # Log traceback for debugging.
        if progress_callback:
            progress_callback(
                -1, f"Error during conversion: {e}"
            )  # -1 indicates error.
        # It's important to let the user know that Word/LibreOffice might be needed.
        detailed_error_message = (
            f"{error_message}. This often occurs if Microsoft Word (Windows) or "
            "LibreOffice (Linux/macOS) is not installed or not found in the system PATH. "
            "Please ensure one of these is installed and accessible."
        )
        raise RuntimeError(detailed_error_message) from e


# Example usage (primarily for testing the module directly)
if __name__ == "__main__":
    # This block runs if the script is executed directly.
    # It's useful for testing the converter in isolation.

    # Configure basic logging for this example.
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(name)s - %(message)s",
    )
    logger.info("Running DOCX to PDF converter example...")

    # Define paths for test input and output files.
    # For this to work, you'd need a sample .docx file.
    # You can create a simple "test_input.docx" manually.
    example_dir = Path.cwd() / "example_conversion_files"
    example_dir.mkdir(exist_ok=True)  # Create if it doesn't exist

    test_input_docx_path = example_dir / "test_input.docx"
    # Create a dummy DOCX file if it doesn't exist and python-docx is available
    try:
        from docx import Document

        if not test_input_docx_path.exists():
            doc = Document()
            doc.add_heading("Test Document for DOCX to PDF", 0)
            doc.add_paragraph("This is a test paragraph.")
            doc.add_paragraph("Another paragraph with some ")
            p = doc.paragraphs[-1]
            p.add_run("bold").bold = True
            p.add_run(" and some ")
            p.add_run("italic.").italic = True
            doc.save(str(test_input_docx_path))
            logger.info(f"Created dummy test DOCX file: {test_input_docx_path}")
        else:
            logger.info(f"Using existing test DOCX file: {test_input_docx_path}")

    except ImportError:
        logger.warning(
            "python-docx is not installed. Cannot create a dummy DOCX file for testing. "
            f"Please create '{test_input_docx_path}' manually to run the example."
        )
        # If no file, exit example to avoid error
        if not test_input_docx_path.exists():
            sys.exit(1)  # Use sys.exit in __main__ block
    except Exception as ex:
        logger.error(f"Could not create dummy DOCX: {ex}")
        if not test_input_docx_path.exists():
            sys.exit(1)

    # Define a simple progress callback for the example.
    def sample_progress_callback(percentage: int, message: str):
        print(f"Progress: {percentage}% - {message}")

    try:
        logger.info(f"Attempting conversion for: {test_input_docx_path}")
        # Call the converter function with an explicit output path
        explicit_output_path = example_dir / "test_output_from_docx.pdf"
        converted_file = convert_docx_to_pdf(
            input_path=test_input_docx_path,
            output_path=explicit_output_path,
            progress_callback=sample_progress_callback,
        )
        logger.info(f"Conversion successful! Output (explicit path): {converted_file}")

        # Call the converter function for auto-generated output path
        logger.info("\nAttempting conversion with auto-generated output path...")
        converted_file_auto = convert_docx_to_pdf(
            input_path=test_input_docx_path,
            # output_path is None here
            progress_callback=sample_progress_callback,
        )
        logger.info(f"Conversion successful! Output (auto path): {converted_file_auto}")

    except ImportError as ie:
        logger.error(
            f"ImportError: {ie}. `docx2pdf` or its dependencies might not be installed."
        )
    except FileNotFoundError as fnfe:
        logger.error(f"FileNotFoundError: {fnfe}. Input file likely missing.")
    except RuntimeError as rte:
        logger.error(f"RuntimeError during conversion: {rte}")
    except Exception as ex:  # Catch any other unexpected error
        logger.error(
            f"An unexpected error occurred in the example: {ex}", exc_info=True
        )
    finally:
        # You might want to keep the generated PDF for inspection
        logger.info("DOCX to PDF example finished.")
        # Example: Clean up specific files
        # if explicit_output_path.exists():
        #     explicit_output_path.unlink()
        # if 'converted_file_auto' in locals() and converted_file_auto.exists():
        #     converted_file_auto.unlink()
        # Keep test_input_docx_path for subsequent runs unless it was auto-generated and you want to clean it.
